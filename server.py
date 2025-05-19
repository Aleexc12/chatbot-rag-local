import os
import time
import uuid
import gc
import json
from flask import Flask, request, jsonify, render_template, send_from_directory, Response
import torch
import fitz
import docx
from werkzeug.utils import secure_filename
from sentence_transformers import CrossEncoder
from qdrant_client import QdrantClient, models
from qdrant_client.http.models import CollectionStatus
import ollama
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument
from dotenv import load_dotenv
import traceback
from qdrant_client.http.models import OptimizersConfigDiff, AliasOperations, CreateAliasOperation
import psycopg2
from psycopg2.extras import DictCursor

load_dotenv()

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'md', 'docx'}
QDRANT_COLLECTION_NAME = "rag_collection_ollama_persistent_v1"

QDRANT_HOST = os.getenv("QDRANT_HOST", "localhost")
QDRANT_PORT = int(os.getenv("QDRANT_PORT", 6333))
QDRANT_API_KEY = os.getenv("QDRANT_API_KEY", None)

PG_HOST = os.getenv("PG_HOST", "localhost")
PG_PORT = os.getenv("PG_PORT", "5432")
PG_USER = os.getenv("PG_USER")
PG_PASSWORD = os.getenv("PG_PASSWORD")
PG_DBNAME = os.getenv("PG_DBNAME")

EMBEDDING_MODEL_NAME_OLLAMA = 'nomic-embed-text'
OLLAMA_EMBEDDING_DIMENSION_DEFAULT = 768

AVAILABLE_LLMS = {
    # "llama3.2:1b": {
    #     "model_name": "llama3.2:1b",
    #     "max_tokens": 2048,
    #     "display_name": "Llama3.2 (1B)"
    # },
    "phi4-mini": {
        "model_name": "phi4-mini",
        "max_tokens": 4096,
        "display_name": "Phi4 Mini (3.8B)"
    },
    # "phi4-mini-reasoning": {
    #     "model_name": "phi4-mini-reasoning",
    #     "max_tokens": 8192,
    #     "display_name": "Phi4 Mini-Reasoning"
    # },
    # "phi4-reasoning": {
    #     "model_name": "phi4-reasoning",
    #     "max_tokens": 16384,
    #     "display_name": "Phi4 Reasoning"
    # },
    # "qwen3:8b": {
    #     "model_name": "qwen3:8b",
    #     "max_tokens": 8192,
    #     "display_name": "Qwen3 (8B)"
    # },
    # "qwen3:14b": {
    #     "model_name": "qwen3:14b",
    #     "max_tokens": 16384,
    #     "display_name": "Qwen3 (14B)"
    # },
}

DEFAULT_LLM_KEY = "phi4-mini"
ROUTER_LLM_KEY = "phi4-mini"

CROSS_ENCODER_MODEL_NAME = 'cross-encoder/ms-marco-MiniLM-L-6-v2'

RECURSIVE_CHUNK_SIZE = 500
RECURSIVE_CHUNK_OVERLAP = 50
INITIAL_QDRANT_SEARCH_LIMIT = 10
FINAL_CONTEXT_CHUNKS_LIMIT = 5
MAX_HISTORY_TURNS = 5
ROUTER_HISTORY_TURNS = 2

os.environ["TF_ENABLE_ONEDNN_OPTS"] = "0"
app = Flask(__name__, template_folder="templates", static_folder='static')
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER): os.makedirs(UPLOAD_FOLDER)
if not os.path.exists('imgs'): os.makedirs('imgs')
if not os.path.exists('static'): os.makedirs('static')

cross_encoder_model = None
qdrant_client = None
text_splitter = None

def get_pg_connection():
    try:
        conn = psycopg2.connect(
            host=PG_HOST,
            port=PG_PORT,
            user=PG_USER,
            password=PG_PASSWORD,
            dbname=PG_DBNAME
        )
        return conn
    except psycopg2.Error as e:
        print(f"Error connecting to PostgreSQL: {e}")
        traceback.print_exc()
        return None

def create_chat_history_table():
    conn = get_pg_connection()
    if not conn:
        print("FATAL: Could not connect to PostgreSQL to create table. Exiting.")
        exit()
    try:
        with conn.cursor() as cur:
            cur.execute("""
                CREATE TABLE IF NOT EXISTS chat_history (
                    id SERIAL PRIMARY KEY,
                    session_id VARCHAR(36) NOT NULL,
                    role VARCHAR(10) NOT NULL CHECK (role IN ('user', 'assistant', 'system')),
                    content TEXT NOT NULL,
                    created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP
                );
            """)
            cur.execute("CREATE INDEX IF NOT EXISTS idx_session_id_created_at ON chat_history (session_id, created_at);")
            conn.commit()
        print("PostgreSQL 'chat_history' table checked/created successfully.")
    except psycopg2.Error as e:
        print(f"Error creating PostgreSQL table: {e}")
        traceback.print_exc()
        conn.rollback()
        print("FATAL: Error creating chat_history table. Exiting.")
        exit()
    finally:
        if conn:
            conn.close()

def get_conversation_history_from_db(session_id, limit_turns=MAX_HISTORY_TURNS):
    history = []
    conn = get_pg_connection()
    if not conn: return history

    messages_to_fetch = (limit_turns * 2) + 1

    try:
        with conn.cursor(cursor_factory=DictCursor) as cur:
            cur.execute("""
                SELECT role, content
                FROM chat_history
                WHERE session_id = %s
                ORDER BY created_at DESC
                LIMIT %s
            """, (session_id, messages_to_fetch))
            raw_messages = cur.fetchall()

        if not raw_messages:
            return history

        reordered_messages = [{"role": msg["role"], "content": msg["content"]} for msg in reversed(raw_messages)]

        system_prompt_message = None
        if reordered_messages and reordered_messages[0]['role'] == 'system':
            system_prompt_message = reordered_messages.pop(0)

        if len(reordered_messages) > limit_turns * 2:
            history = reordered_messages[-(limit_turns * 2):]
        else:
            history = reordered_messages

        if system_prompt_message:
            history.insert(0, system_prompt_message)

    except psycopg2.Error as e:
        print(f"Error fetching conversation history from DB for session {session_id}: {e}")
        traceback.print_exc()
    finally:
        if conn:
            conn.close()
    return history

def add_message_to_db(session_id, role, content):
    conn = get_pg_connection()
    if not conn: return False
    try:
        with conn.cursor() as cur:
            cur.execute("""
                INSERT INTO chat_history (session_id, role, content)
                VALUES (%s, %s, %s)
            """, (session_id, role, content))
            conn.commit()
        return True
    except psycopg2.Error as e:
        print(f"Error adding message to DB for session {session_id}: {e}")
        traceback.print_exc()
        conn.rollback()
        return False
    finally:
        if conn:
            conn.close()

def get_all_messages_for_session_from_db(session_id):
    messages = []
    conn = get_pg_connection()
    if not conn: return messages

    try:
        with conn.cursor(cursor_factory=DictCursor) as cur:
            cur.execute("""
                SELECT role, content, TO_CHAR(created_at AT TIME ZONE 'UTC', 'YYYY-MM-DD HH24:MI:SS TZ') as timestamp
                FROM chat_history
                WHERE session_id = %s
                ORDER BY created_at ASC
            """, (session_id,))
            raw_messages = cur.fetchall()
        messages = [{"role": msg["role"], "content": msg["content"], "timestamp": msg["timestamp"]} for msg in raw_messages]
    except psycopg2.Error as e:
        print(f"Error fetching all messages from DB for session {session_id}: {e}")
        traceback.print_exc()
    finally:
        if conn:
            conn.close()
    return messages

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def check_and_pull_ollama_model(model_name_tag):
    try:
        print(f"Checking for Ollama model: {model_name_tag}...")
        ollama_list_response = ollama.list()
        is_available = any(
            model_info['name'] == model_name_tag
            for model_info in ollama_list_response.get('models', [])
            if isinstance(model_info, dict) and 'name' in model_info
        )
        if not is_available:
            print(f"Ollama model '{model_name_tag}' not found locally. Pulling...")
            ollama.pull(model_name_tag)
            print(f"Successfully pulled '{model_name_tag}'.")
            return True
        else:
            print(f"Ollama model '{model_name_tag}' is available locally.")
            return True
    except Exception as e:
        print(f"Error in check_and_pull_ollama_model for '{model_name_tag}': {e}")
        print(f"Ensure Ollama is running and model tag is correct. Try: `ollama pull {model_name_tag}` manually.")
        return False

def initialize_dependencies():
    global cross_encoder_model, qdrant_client, text_splitter
    print("--- Initializing Dependencies ---")

    print("Initializing PostgreSQL connection and schema...")
    if not all([PG_HOST, PG_PORT, PG_USER, PG_PASSWORD, PG_DBNAME]):
        print("FATAL: PostgreSQL environment variables (PG_HOST, PG_PORT, PG_USER, PG_PASSWORD, PG_DBNAME) are not fully set. Exiting.")
        exit()
    create_chat_history_table()
    conn_test = get_pg_connection()
    if conn_test:
        print("PostgreSQL connection successful.")
        conn_test.close()
    else:
        print("FATAL: PostgreSQL connection test failed during initialization. Exiting.")
        exit()

    try:
        print("Checking Ollama service...")
        ollama.list()
        print("Ollama service detected.")

        print("--- Checking availability of configured LLM models ---")
        if ROUTER_LLM_KEY not in AVAILABLE_LLMS:
            print(f"FATAL: ROUTER_LLM_KEY '{ROUTER_LLM_KEY}' is not defined in AVAILABLE_LLMS. Exiting.")
            exit()

        for llm_key, config in AVAILABLE_LLMS.items():
            model_tag_to_check = config["model_name"]
            print(f"Verifying LLM: {config['display_name']} (Ollama tag: {model_tag_to_check})")
            if not check_and_pull_ollama_model(model_tag_to_check):
                print(f"WARNING: LLM '{model_tag_to_check}' for '{llm_key}' could not be verified or pulled. It may not be usable.")
        print("--- LLM model check complete ---")

        print(f"Verifying Embedding Model: {EMBEDDING_MODEL_NAME_OLLAMA}")
        if not check_and_pull_ollama_model(EMBEDDING_MODEL_NAME_OLLAMA):
            print(f"FATAL: Could not ensure embedding model '{EMBEDDING_MODEL_NAME_OLLAMA}' is available. Exiting.")
            exit()

        ce_device = 'cuda' if torch.cuda.is_available() else 'cpu'
        print(f"Loading Cross-encoder '{CROSS_ENCODER_MODEL_NAME}' on {ce_device}...")
        cross_encoder_model = CrossEncoder(CROSS_ENCODER_MODEL_NAME, device=ce_device, max_length=512)
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=RECURSIVE_CHUNK_SIZE,
            chunk_overlap=RECURSIVE_CHUNK_OVERLAP,
            length_function=len,
            is_separator_regex=False,
        )
        print("Core text processing dependencies loaded.")
    except Exception as e:
        print(f"FATAL: Failed to load/initialize core models: {e}")
        traceback.print_exc()
        exit()

    try:
        print(f"Initializing Qdrant client for persistent storage at {QDRANT_HOST}:{QDRANT_PORT}...")
        if QDRANT_API_KEY:
            qdrant_client = QdrantClient(host=QDRANT_HOST, port=QDRANT_PORT, api_key=QDRANT_API_KEY, timeout=20)
        else:
            qdrant_client = QdrantClient(host=QDRANT_HOST, port=QDRANT_PORT, timeout=20)

        print("Pinging Qdrant service...")
        collections_response = qdrant_client.get_collections()
        print(f"Successfully connected to Qdrant service. Found {len(collections_response.collections)} existing collections.")
        
        emb_dim = OLLAMA_EMBEDDING_DIMENSION_DEFAULT
        try:
            print(f"Getting embedding dimension for '{EMBEDDING_MODEL_NAME_OLLAMA}'...")
            response = ollama.embeddings(model=EMBEDDING_MODEL_NAME_OLLAMA, prompt="hello")
            if "embedding" in response and isinstance(response["embedding"], list) and len(response["embedding"]) > 0:
                emb_dim = len(response["embedding"])
                print(f"Dynamically determined embedding dimension: {emb_dim}")
            else:
                print(f"Warning: Could not get valid embedding dim. Using default {OLLAMA_EMBEDDING_DIMENSION_DEFAULT}.")
        except Exception as e_emb_dim:
            print(f"Warning: Could not get embedding dim, using default {OLLAMA_EMBEDDING_DIMENSION_DEFAULT}. Error: {e_emb_dim}")
        
        collection_exists = False
        try:
            collection_info = qdrant_client.get_collection(collection_name=QDRANT_COLLECTION_NAME)
            print(f"Found existing Qdrant collection '{QDRANT_COLLECTION_NAME}' with {collection_info.points_count} points.")
            if collection_info.config.params.vectors.size != emb_dim:
                print(f"WARNING: Existing collection '{QDRANT_COLLECTION_NAME}' has vector dimension "
                      f"{collection_info.config.params.vectors.size}, but embedding model expects {emb_dim}. "
                      f"This will likely cause errors. Consider deleting and recreating the collection manually in Qdrant "
                      f"or update EMBEDDING_MODEL_NAME_OLLAMA and its dimension.")
            collection_exists = True
        except Exception as e:
            error_message = str(e).lower()
            if "not found" in error_message or "status_code=404" in error_message or "collectionnotfoundexception" in error_message:
                print(f"Qdrant collection '{QDRANT_COLLECTION_NAME}' not found. Creating new collection...")
            elif "connection refused" in error_message or "failed to connect" in error_message:
                print(f"FATAL: Could not connect to Qdrant at {QDRANT_HOST}:{QDRANT_PORT}. Ensure Qdrant service is running and accessible. Error: {e}")
                traceback.print_exc()
                exit()
            else:
                print(f"Error checking for collection '{QDRANT_COLLECTION_NAME}': {e}. Will attempt to create.")
        
        if not collection_exists:
            try:
                print(f"Attempting to create Qdrant collection '{QDRANT_COLLECTION_NAME}' with dimension {emb_dim}...")
                qdrant_client.create_collection(
                    collection_name=QDRANT_COLLECTION_NAME,
                    vectors_config=models.VectorParams(size=emb_dim, distance=models.Distance.COSINE)
                )
                print(f"Qdrant collection '{QDRANT_COLLECTION_NAME}' created successfully.")
            except Exception as e_create:
                print(f"FATAL: Failed to create Qdrant collection '{QDRANT_COLLECTION_NAME}': {e_create}")
                print("This could be due to a connection issue, incompatible parameters, or Qdrant server problems.")
                traceback.print_exc()
                exit()
        print(f"Qdrant setup for collection '{QDRANT_COLLECTION_NAME}' (dim: {emb_dim}) complete.")
        print("--- All Dependencies Initialized Successfully ---")
    except Exception as e:
        print(f"FATAL: Failed to initialize or connect to Qdrant: {e}")
        print("Ensure Qdrant service is running and accessible (e.g., `docker ps` should show your qdrant container).")
        print(f"Attempted to connect to Qdrant at: {QDRANT_HOST}:{QDRANT_PORT}")
        traceback.print_exc()
        exit()

initialize_dependencies()

@app.route('/static/<path:filename>')
def serve_static(filename): return send_from_directory('static', filename)

@app.route('/imgs/<path:filename>')
def serve_image(filename): return send_from_directory('imgs', filename)

@app.route("/")
def home():
    session_id = str(uuid.uuid4())
    print(f"Home page. Initial session_id for new interaction: {session_id}.")
    return render_template("index.html",
                           current_session_id=session_id,
                           available_llms=AVAILABLE_LLMS,
                           default_llm_key=DEFAULT_LLM_KEY)

@app.route('/get_chat_sessions', methods=['GET'])
def get_chat_sessions_route():
    sessions_data = []
    conn = get_pg_connection()
    if not conn:
        return jsonify({"error": "Could not connect to database"}), 500
    try:
        with conn.cursor(cursor_factory=DictCursor) as cur:
            cur.execute("""
                SELECT
                    s.session_id,
                    MAX(s.created_at) as last_activity_at,
                    (SELECT content FROM chat_history 
                     WHERE session_id = s.session_id AND role = 'user' 
                     ORDER BY created_at ASC LIMIT 1) as first_user_message
                FROM chat_history s
                GROUP BY s.session_id
                ORDER BY last_activity_at DESC;
            """)
            sessions = cur.fetchall()
            for session in sessions:
                display_name = "Chat " + session['session_id'][:8]
                if session['first_user_message']:
                    words = session['first_user_message'].split()
                    display_name = ' '.join(words[:5])
                    if len(words) > 5:
                        display_name += "..."
                
                sessions_data.append({
                    "session_id": session['session_id'],
                    "display_name": display_name,
                    "last_activity": session['last_activity_at'].isoformat() if session['last_activity_at'] else None
                })
    except psycopg2.Error as e:
        print(f"Error fetching chat sessions: {e}")
        traceback.print_exc()
        return jsonify({"error": "Error fetching sessions"}), 500
    finally:
        if conn:
            conn.close()
    return jsonify(sessions_data)

@app.route('/get_messages_for_session/<session_id>', methods=['GET'])
def get_messages_for_session_route(session_id):
    if not session_id:
        return jsonify({"error": "Session ID is required"}), 400
    
    messages = get_all_messages_for_session_from_db(session_id)
    return jsonify(messages)

@app.route("/chat", methods=["POST"])
def chat():
    data = request.get_json()
    user_message_text = data.get("message", "").strip()
    selected_llm_key = data.get("model", DEFAULT_LLM_KEY)
    session_id = data.get("session_id", None)

    if not user_message_text:
        return jsonify({"error": "Empty message received."}), 400
    
    if not session_id:
        return jsonify({"error": "Session ID is missing."}), 400

    if selected_llm_key not in AVAILABLE_LLMS:
        print(f"Warning: Selected LLM key '{selected_llm_key}' not in AVAILABLE_LLMS. Falling back to default.")
        selected_llm_key = DEFAULT_LLM_KEY
    
    current_llm_config = AVAILABLE_LLMS[selected_llm_key]
    current_llm_name_ollama = current_llm_config["model_name"]
    current_llm_max_tokens = current_llm_config["max_tokens"]
    max_context_chars_for_llm = int((current_llm_max_tokens * 0.6) * 3.5)

    router_llm_ollama_tag = AVAILABLE_LLMS[ROUTER_LLM_KEY]["model_name"]
    print(f"Chat request. User LLM: {current_llm_config['display_name']}, Router LLM: {AVAILABLE_LLMS[ROUTER_LLM_KEY]['display_name']}, Session: {session_id}, Streaming: True")

    current_conversation_history = get_conversation_history_from_db(session_id, limit_turns=MAX_HISTORY_TURNS)
    
    is_new_conversation_in_db = not bool(get_conversation_history_from_db(session_id, limit_turns=1))
    if is_new_conversation_in_db:
        system_prompt_content_default = "You are a helpful and concise AI assistant."
        add_message_to_db(session_id, "system", system_prompt_content_default)
        if not any(m['role'] == 'system' for m in current_conversation_history):
            current_conversation_history.insert(0, {"role": "system", "content": system_prompt_content_default})
        print(f"First message for session {session_id}. Added default system prompt to DB and history.")

    action_to_take = "direct_answer"
    router_history_ollama = get_conversation_history_from_db(session_id, limit_turns=ROUTER_HISTORY_TURNS)
    
    router_user_prompt_content = (
        f"[Routing Task] User Question: \"{user_message_text}\"\n"
        "Based on the User Question, do you need to search uploaded documents for an answer?\n"
        "Respond with ONLY ONE of these words: \"rag_search\" or \"direct_answer\".\n"
        "Examples:\n"
        "User Question: \"summarize the uploaded PDF about llamas\"\nYour decision: rag_search\n"
        "User Question: \"what is the capital of France?\"\nYour decision: direct_answer\n"
        "User Question: \"tell me more about the project details in the report\"\nYour decision: rag_search\n"
        "Your decision:"
    )
    router_messages_ollama = router_history_ollama + [{"role": "user", "content": router_user_prompt_content}]
    
    router_system_prompt = "You are a routing agent. Decide 'rag_search' or 'direct_answer'."
    if not any(m['role'] == 'system' for m in router_messages_ollama):
        router_messages_ollama.insert(0, {"role": "system", "content": router_system_prompt})
    elif router_messages_ollama[0]['role'] != 'system':
         router_messages_ollama.insert(0, {"role": "system", "content": router_system_prompt})
    else: 
        router_messages_ollama[0]['content'] = router_system_prompt

    try:
        router_response = ollama.chat(
            model=router_llm_ollama_tag, 
            messages=router_messages_ollama, 
            stream=False, 
            options={"temperature": 0.0, "num_predict": 15}
        )
        decision_raw = router_response['message']['content'].strip()
        decision_clean = ''.join(filter(str.isalnum, decision_raw.lower().replace("_", "")))
        print(f"Router Raw: '{decision_raw}' -> Clean: '{decision_clean}' (using {router_llm_ollama_tag})")
        
        if "ragsearch" in decision_clean: action_to_take = "rag_search"
        elif "directanswer" in decision_clean: action_to_take = "direct_answer"
        else:
            keywords_for_rag = ["document", "pdf", "file", "summary", "summarize", "detail", "information", "content", "say about", "uploaded", "context", "retrieve", "chapter", "section", "article"]
            if any(keyword in user_message_text.lower() for keyword in keywords_for_rag): action_to_take = "rag_search"
            else: action_to_take = "direct_answer"
            print(f"Router decision unclear ('{decision_raw}'). Fallback based on keywords: {action_to_take.upper()}.")
    except Exception as e_router:
        print(f"Router Error with Ollama ({router_llm_ollama_tag}): {e_router}. Fallback to keyword matching.")
        traceback.print_exc()
        keywords_for_rag = ["document", "pdf", "file", "summary", "summarize", "detail", "information", "content", "say about", "uploaded", "context", "retrieve", "chapter", "section", "article"]
        if any(keyword in user_message_text.lower() for keyword in keywords_for_rag): action_to_take = "rag_search"
        else: action_to_take = "direct_answer"
        print(f"Fallback (on router error): {action_to_take.upper()}.")
    print(f"==> Router Decision: {action_to_take.upper()}")

    context_prompt_insert = ""
    retrieved_sources = []
    if action_to_take == "rag_search":
        print(f"==> Performing RAG Search (Max context chars for LLM '{current_llm_name_ollama}': {max_context_chars_for_llm})...")
        try:
            response_emb = ollama.embeddings(model=EMBEDDING_MODEL_NAME_OLLAMA, prompt=user_message_text)
            query_emb = response_emb['embedding']
            hits = qdrant_client.search(
                collection_name=QDRANT_COLLECTION_NAME,
                query_vector=query_emb,
                limit=INITIAL_QDRANT_SEARCH_LIMIT
            )
            final_hits_for_context = []
            if hits:
                if cross_encoder_model:
                    pairs_for_reranking = [[user_message_text, hit.payload['text']] for hit in hits if hit.payload and 'text' in hit.payload]
                    if pairs_for_reranking:
                        scores = cross_encoder_model.predict(pairs_for_reranking, show_progress_bar=False)
                        scored_hits = sorted(zip(scores, hits), key=lambda x: x[0], reverse=True)
                        final_hits_for_context = [hit for score, hit in scored_hits[:FINAL_CONTEXT_CHUNKS_LIMIT]]
                    else:
                        print("RAG: No valid pairs for cross-encoder reranking, using initial Qdrant hits.")
                        final_hits_for_context = hits[:FINAL_CONTEXT_CHUNKS_LIMIT]
                else:
                    final_hits_for_context = hits[:FINAL_CONTEXT_CHUNKS_LIMIT]
            
            if final_hits_for_context:
                context_parts = []; current_chars_in_context = 0
                for i, hit in enumerate(final_hits_for_context):
                    text_content = hit.payload['text']
                    source_filename = hit.payload.get('source', 'Unknown source')
                    if (current_chars_in_context + len(text_content) + len(f"Context Snippet {i+1} from '{source_filename}':\n\n\n---\n\n")) <= max_context_chars_for_llm:
                        context_parts.append(f"Context Snippet {i+1} from '{source_filename}':\n{text_content}")
                        retrieved_sources.append({"filename": source_filename, "id": str(hit.id), "text": text_content[:100] + "..."})
                        current_chars_in_context += len(text_content) + len(f"Context Snippet {i+1} from '{source_filename}':\n\n\n---\n\n")
                    else: 
                        print(f"RAG: Context limit ({max_context_chars_for_llm} chars) reached. Stopped adding snippets.")
                        break
                if context_parts:
                    context_prompt_insert = ("Based ONLY on the following text snippets from uploaded documents, answer the user's question. "
                                             "If the answer is not in the provided snippets, say so. Do not use any other knowledge.\n\n"
                                             "---\n\n".join(context_parts) + "\n\n---\n\n")
                else: 
                    action_to_take = "direct_answer"
                    print("RAG: No usable context snippets after filtering by length. Switching to direct_answer.")
            else: 
                action_to_take = "direct_answer"
                print("RAG: No matching documents found after search/reranking. Switching to direct_answer.")
        except Exception as e_rag:
            print(f"RAG Processing Error: {e_rag}"); traceback.print_exc()
            action_to_take = "direct_answer"
            print("RAG: Error during RAG process. Switching to direct_answer.")

    final_ollama_messages = current_conversation_history.copy()

    user_turn_content = user_message_text
    if action_to_take == "rag_search" and context_prompt_insert:
        user_turn_content = f"{context_prompt_insert}User Question: {user_message_text}"
    
    final_ollama_messages.append({"role": "user", "content": user_turn_content})

    accumulated_response_text = ""

    def generate_chat_responses():
        nonlocal accumulated_response_text
        accumulated_response_text = ""
        try:
            print(f"==> Streaming response via Ollama '{current_llm_name_ollama}' (Strategy: {action_to_take.upper()}).")
            generation_options = {
                "temperature": 0.2 if action_to_take == "rag_search" and context_prompt_insert else 0.7, 
                "top_p": 0.9
            }
            stream = ollama.chat(
                model=current_llm_name_ollama,
                messages=final_ollama_messages, 
                stream=True, 
                options=generation_options
            )

            for chunk in stream:
                if chunk['done'] is False:
                    content_piece = chunk['message']['content']
                    accumulated_response_text += content_piece
                    sse_event = {"type": "content", "text": content_piece}
                    yield f"data: {json.dumps(sse_event)}\n\n"
                else: 
                    print(f"Ollama Stream 'done'. Full response length for '{current_llm_name_ollama}': {len(accumulated_response_text)}")
                    pass 

            final_event_data = {
                "type": "final",
                "full_text": accumulated_response_text,
                "sources": retrieved_sources if action_to_take == "rag_search" and retrieved_sources else [],
            }
            yield f"data: {json.dumps(final_event_data)}\n\n"

            add_message_to_db(session_id, "user", user_message_text)
            add_message_to_db(session_id, "assistant", accumulated_response_text)
            
        except ollama.ResponseError as e_ollama_resp:
            error_message = f"Ollama API error during generation with '{current_llm_name_ollama}': {e_ollama_resp.error} (Status: {e_ollama_resp.status_code})"
            print(f"CRITICAL OLLAMA RESPONSE ERROR (stream): {error_message}")
            error_event = {"type": "error", "message": error_message}
            yield f"data: {json.dumps(error_event)}\n\n"
        except Exception as e_generation:
            error_message = f"Server error during generation with '{current_llm_name_ollama}': {str(e_generation)}"
            print(f"General Generation Error (stream): {error_message}"); traceback.print_exc()
            error_event = {"type": "error", "message": error_message}
            yield f"data: {json.dumps(error_event)}\n\n"

    return Response(generate_chat_responses(), mimetype='text/event-stream')

@app.route('/upload', methods=['POST'])
def upload_file_route():
    global text_splitter
    if 'files[]' not in request.files:
        return jsonify({"error": "No file part"}), 400
    
    files_from_request = request.files.getlist('files[]')
    if not files_from_request or all(not f.filename for f in files_from_request):
        return jsonify({"error": "No files selected"}), 400

    results_summary = []; temp_saved_paths = {}
    for file_obj in files_from_request:
        original_filename = file_obj.filename
        current_file_result = {"original_filename": original_filename, "status": "error", "message": "Skipped", "chunks_indexed": 0}

        if original_filename and allowed_file(original_filename):
            secure_fn = secure_filename(original_filename)
            temp_file_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_fn)
            temp_saved_paths[original_filename] = temp_file_path
            
            try:
                file_obj.save(temp_file_path)
                extracted_text = ""
                file_ext_lower = original_filename.rsplit('.', 1)[-1].lower()

                if file_ext_lower == 'pdf':
                    with fitz.open(temp_file_path) as doc:
                        extracted_text = "\n\n".join([page.get_text("text", sort=True).strip() for page in doc if page.get_text("text", sort=True).strip()])
                elif file_ext_lower == 'docx':
                    doc = docx.Document(temp_file_path)
                    extracted_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
                elif file_ext_lower in ('txt', 'md'):
                    try: 
                        with open(temp_file_path, 'r', encoding='utf-8') as f: extracted_text = f.read()
                    except UnicodeDecodeError: 
                        with open(temp_file_path, 'r', encoding='latin-1') as f: extracted_text = f.read()
                
                if not extracted_text.strip(): 
                    raise ValueError("No text extracted or file is empty.")

                langchain_docs_to_split = [LangchainDocument(page_content=extracted_text, metadata={"source": original_filename})]
                if not text_splitter:
                    raise RuntimeError("TextSplitter not initialized.")
                split_chunks_as_docs = text_splitter.split_documents(langchain_docs_to_split)
                chunks_of_text = [chunk.page_content for chunk in split_chunks_as_docs if chunk.page_content.strip()]

                if not chunks_of_text: 
                    raise ValueError("No processable text chunks after splitting.")
                
                print(f"Processing {len(chunks_of_text)} chunks for '{original_filename}' using embedding model '{EMBEDDING_MODEL_NAME_OLLAMA}'...")
                
                chunk_embeddings = []
                for i, chunk_text_content in enumerate(chunks_of_text):
                    try:
                        embedding_response = ollama.embeddings(model=EMBEDDING_MODEL_NAME_OLLAMA, prompt=chunk_text_content)
                        chunk_embeddings.append(embedding_response['embedding'])
                    except Exception as e_emb_chunk:
                        print(f"Error embedding chunk {i+1}/{len(chunks_of_text)} for '{original_filename}': {e_emb_chunk}. Skipping this chunk.")
                        chunk_embeddings.append(None)
                
                valid_embeddings_data = []
                for i, emb_vector in enumerate(chunk_embeddings):
                    if emb_vector:
                        valid_embeddings_data.append((emb_vector, chunks_of_text[i]))
                
                if not valid_embeddings_data: 
                    raise ValueError("No embeddings were successfully generated for any chunks.")

                points_to_upsert_in_qdrant = [
                    models.PointStruct(
                        id=str(uuid.uuid4()), 
                        vector=emb_vector, 
                        payload={"text": chunk_text_payload, "source": original_filename}
                    )
                    for emb_vector, chunk_text_payload in valid_embeddings_data
                ]
                
                if points_to_upsert_in_qdrant:
                    qdrant_client.upsert(
                        collection_name=QDRANT_COLLECTION_NAME, 
                        points=points_to_upsert_in_qdrant, 
                        wait=True
                    )
                    current_file_result["status"] = "success"
                    current_file_result["message"] = f"Processed and indexed {len(points_to_upsert_in_qdrant)} chunks."
                    current_file_result["chunks_indexed"] = len(points_to_upsert_in_qdrant)
                else:
                    current_file_result["message"] = "No valid text chunks with embeddings to index."
            
            except Exception as e_proc_file:
                current_file_result["message"] = f"Processing failed: {str(e_proc_file)}"
                print(f"Error processing file '{original_filename}': {e_proc_file}"); traceback.print_exc()
        else:
            current_file_result["message"] = "File type not allowed or filename invalid."
        
        results_summary.append(current_file_result)

    for path_to_remove in temp_saved_paths.values():
        if os.path.exists(path_to_remove):
            try: 
                os.remove(path_to_remove)
                print(f"Cleaned up temporary file: '{path_to_remove}'")
            except Exception as e_cleanup: 
                print(f"Warning: Failed to clean up temporary file '{path_to_remove}': {e_cleanup}")
    
    if not results_summary:
        return jsonify({"error": "No files were processed."}), 400
    
    all_successful = all(r['status'] == 'success' for r in results_summary)
    any_successful = any(r['status'] == 'success' for r in results_summary)
    
    if all_successful:
        http_status_code = 200
    elif any_successful:
        http_status_code = 207
    else:
        http_status_code = 400

    return jsonify({"results": results_summary}), http_status_code

@app.route('/switch_model', methods=['POST'])
def switch_model_api_route():
    data = request.get_json()
    model_key_from_js = data.get('model')

    if not model_key_from_js or model_key_from_js not in AVAILABLE_LLMS:
        current_default_model_name = AVAILABLE_LLMS.get(DEFAULT_LLM_KEY, {}).get("display_name", "the default model")
        return jsonify({"status": "error", "message": f"Invalid model key. Using {current_default_model_name}.", "loaded_model_key": DEFAULT_LLM_KEY}), 400

    ollama_tag_to_verify = AVAILABLE_LLMS[model_key_from_js]["model_name"]
    display_name_to_verify = AVAILABLE_LLMS[model_key_from_js]["display_name"]
    
    try:
        print(f"Verifying model '{ollama_tag_to_verify}' ({display_name_to_verify}) for switch...")
        if check_and_pull_ollama_model(ollama_tag_to_verify):
            print(f"Model '{ollama_tag_to_verify}' confirmed available for use.")
            return jsonify({
                "status": "success", 
                "message": f"Switched to model: {display_name_to_verify}", 
                "loaded_model_key": model_key_from_js
            })
        else:
            current_default_model_name = AVAILABLE_LLMS[DEFAULT_LLM_KEY]["display_name"]
            return jsonify({
                "status": "error", 
                "message": f"Failed to verify or pull model '{display_name_to_verify}'. Reverting to {current_default_model_name}.",
                "loaded_model_key": DEFAULT_LLM_KEY 
            }), 500

    except Exception as e_switch:
        print(f"Error verifying model '{ollama_tag_to_verify}' during switch: {e_switch}")
        traceback.print_exc()
        current_default_model_name = AVAILABLE_LLMS[DEFAULT_LLM_KEY]["display_name"]
        return jsonify({
            "status": "error", 
            "message": f"An unexpected error occurred while trying to switch to '{display_name_to_verify}'. Reverting to {current_default_model_name}. Error: {str(e_switch)[:100]}...",
            "loaded_model_key": DEFAULT_LLM_KEY 
        }), 500

if __name__ == "__main__":
    print(f"\n--- Starting Flask Server for RAG Chat Application ---")
    print("IMPORTANT: Ensure your Qdrant service is running and accessible.")
    print(f"This app will try to connect to Qdrant at: {QDRANT_HOST}:{QDRANT_PORT}")
    print(f"Data in Qdrant collection '{QDRANT_COLLECTION_NAME}' WILL PERSIST if Qdrant is run with a persistent volume.")
    print("IMPORTANT: Ensure your PostgreSQL service is running and accessible.")
    print(f"This app will try to connect to PostgreSQL at: {PG_HOST}:{PG_PORT} (DB: {PG_DBNAME})")
    print("Chat history WILL PERSIST in PostgreSQL.")
    print("If the collection/table doesn't exist or its schema is incompatible, it will be (re)created.")
    print(f"Available LLMs (Verify tags with 'ollama list'):")
    for key, config in AVAILABLE_LLMS.items():
        is_default = "(Default)" if key == DEFAULT_LLM_KEY else ""
        is_router = "(Router)" if key == ROUTER_LLM_KEY else ""
        print(f"  - Key: '{key}', Display: \"{config['display_name']}\", Ollama Tag: \"{config['model_name']}\" {is_default} {is_router}")
    print(f"Embedding Model: {EMBEDDING_MODEL_NAME_OLLAMA}")
    print(f"--- --- --- --- --- --- --- --- --- --- --- --- --- ---")
    
    app.run(host="0.0.0.0", port=5001, debug=False)