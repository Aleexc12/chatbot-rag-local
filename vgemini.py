import os
import time
import uuid
import gc
from flask import Flask, request, jsonify, render_template, send_from_directory
import torch
import fitz  # PyMuPDF
import docx
from werkzeug.utils import secure_filename
from sentence_transformers import SentenceTransformer, CrossEncoder
from qdrant_client import QdrantClient, models
# from transformers import AutoModelForCausalLM, AutoTokenizer, BitsAndBytesConfig # Kept if you want to re-add local models
# import accelerate
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument
from dotenv import load_dotenv
import google.generativeai as genai

# --- Configuration ---
load_dotenv() 

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'md', 'docx'}
QDRANT_COLLECTION_NAME = "rag_collection_gemini_v1_nocitechunk" # New name for clarity
EMBEDDING_MODEL_NAME = 'all-MiniLM-L6-v2'
CROSS_ENCODER_MODEL_NAME = 'cross-encoder/ms-marco-MiniLM-L-6-v2' 
MAX_CONTEXT_TOKENS_FOR_LLM = 3000 
RECURSIVE_CHUNK_SIZE = 700 
RECURSIVE_CHUNK_OVERLAP = 70
INITIAL_QDRANT_SEARCH_LIMIT = 10
FINAL_CONTEXT_CHUNKS_LIMIT = 4 

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
GEMINI_ROUTER_MODEL_NAME = "gemini-1.5-flash-latest" 
GEMINI_CHAT_MODEL_NAME = "gemini-1.5-pro-latest" 

if GOOGLE_API_KEY:
    genai.configure(api_key=GOOGLE_API_KEY)
else:
    print("CRITICAL: GOOGLE_API_KEY environment variable not set. Gemini API calls will fail.")

os.environ["TF_ENABLE_ONEDNN_OPTS"] = "0"
app = Flask(__name__, template_folder="templates", static_folder='static')
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER): os.makedirs(UPLOAD_FOLDER)
if not os.path.exists('imgs'): os.makedirs('imgs')
if not os.path.exists('static'): os.makedirs('static')

embedding_model = None
cross_encoder_model = None
qdrant_client = None
conversation_history = [] 
text_splitter = None

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def generate_with_gemini(model_name, prompt_parts_list, temperature=0.6, top_p=0.9, max_output_tokens=1024, is_router_call=False):
    if not GOOGLE_API_KEY:
        return "Error: Gemini API key not configured."
    try:
        model = genai.GenerativeModel(model_name)
        generation_config = genai.types.GenerationConfig(
            temperature=temperature,
            top_p=top_p,
            max_output_tokens=max_output_tokens
        )
        contents_for_gemini = []
        for part in prompt_parts_list:
            role = "model" if part.get("role") == "assistant" else part.get("role", "user")
            content_parts = part.get("parts", [])
            if isinstance(content_parts, str): 
                content_parts = [content_parts]
            elif not isinstance(content_parts, list) or not all(isinstance(p, str) for p in content_parts):
                content_parts = [""] # Fallback for invalid format
            contents_for_gemini.append({"role": role, "parts": content_parts})
        
        if is_router_call and len(contents_for_gemini) == 1 and contents_for_gemini[0]["role"] == "user":
            pass 
        elif is_router_call: 
            last_user_part = next((p for p in reversed(contents_for_gemini) if p["role"] == "user"), None)
            contents_for_gemini = [last_user_part] if last_user_part else [{"role": "user", "parts": [""]}]

        # print(f"--- Sending to Gemini ({model_name}) --- \nContents: {contents_for_gemini if len(contents_for_gemini) < 3 else str(contents_for_gemini[:2]) + '...'} \n---------------------------------")
        
        response = model.generate_content(contents_for_gemini, generation_config=generation_config)
        
        if not response.candidates or not response.candidates[0].content.parts:
            feedback_info = f"Prompt Feedback: {response.prompt_feedback}" if response.prompt_feedback else "No specific feedback."
            print(f"Gemini API ({model_name}): No content in response or response blocked. {feedback_info}")
            # Check for safety ratings if available
            if response.candidates and response.candidates[0].finish_reason.name == "SAFETY":
                print(f"Safety Ratings: {response.candidates[0].safety_ratings}")
            return "Error: No response from Gemini or content was blocked due to safety/other reasons."
        
        return response.text
    except Exception as e:
        print(f"Error calling Gemini API ({model_name}): {e}")
        return f"Error: Could not communicate with Gemini API."

def initialize_dependencies():
    global embedding_model, cross_encoder_model, qdrant_client, text_splitter
    try:
        emb_device = 'cuda' if torch.cuda.is_available() else 'cpu'
        embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME, device=emb_device)
        cross_encoder_model = CrossEncoder(CROSS_ENCODER_MODEL_NAME, device=emb_device)
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=RECURSIVE_CHUNK_SIZE, chunk_overlap=RECURSIVE_CHUNK_OVERLAP)
        print(f"Embedding, CrossEncoder, and TextSplitter models loaded on {emb_device}.")
    except Exception as e: print(f"FATAL: Failed to load ML support models: {e}"); exit()
    try:
        qdrant_client = QdrantClient(":memory:")
        emb_dim = embedding_model.get_sentence_embedding_dimension()
        if not qdrant_client.collection_exists(collection_name=QDRANT_COLLECTION_NAME):
            qdrant_client.create_collection(collection_name=QDRANT_COLLECTION_NAME, vectors_config=models.VectorParams(size=emb_dim, distance=models.Distance.COSINE))
        print(f"Qdrant initialized (Collection: {QDRANT_COLLECTION_NAME}).")
    except Exception as e: print(f"FATAL: Failed to initialize Qdrant: {e}"); exit()

initialize_dependencies()

@app.route('/static/<path:filename>')
def serve_static(filename): return send_from_directory('static', filename)

@app.route('/imgs/<path:filename>')
def serve_image(filename): return send_from_directory('imgs', filename)

@app.route("/")
def home():
    global conversation_history
    conversation_history = []
    try:
        if qdrant_client:
            emb_dim = embedding_model.get_sentence_embedding_dimension()
            qdrant_client.recreate_collection(collection_name=QDRANT_COLLECTION_NAME, vectors_config=models.VectorParams(size=emb_dim, distance=models.Distance.COSINE))
            print(f"Qdrant collection '{QDRANT_COLLECTION_NAME}' reset.")
    except Exception as e: print(f"Error resetting Qdrant: {e}")
    return render_template("index.html")

@app.route("/switch_model", methods=["POST"])
def switch_model_route():
    global conversation_history
    conversation_history = []
    print("Model switch endpoint called. Currently using Gemini. Conversation history cleared.")
    return jsonify({"status": "success", "message": "Currently using Gemini API. History cleared.", "loaded_model": "Gemini"}), 200

@app.route("/chat", methods=["POST"])
def chat():
    global conversation_history
    data = request.get_json()
    user_message_text = data.get("message", "").strip()

    if not user_message_text: return jsonify({"error": "Empty message."}), 400
    if not GOOGLE_API_KEY: return jsonify({"error": "Gemini API key not configured on server."}), 500

    action_to_take = "direct_answer" 
    router_history_text_parts = []
    if len(conversation_history) > 0:
        for msg in conversation_history[-4:]:
            role = "User" if msg["role"] == "user" else "Assistant"
            content = " ".join(msg.get("parts", [""]))
            router_history_text_parts.append(f"{role}: {content}")
    recent_history_str_for_router = "\n".join(router_history_text_parts)
    if recent_history_str_for_router:
        recent_history_str_for_router = f"Previous Conversation Snippet:\n{recent_history_str_for_router}\n"

    router_prompt = f"""You are a routing assistant. Your task is to decide if a user's question requires searching in uploaded documents (RAG) or can be answered directly.
{recent_history_str_for_router}
User Question: "{user_message_text}"

Consider the following actions:
- "rag_search": Use this if the question specifically asks for information likely found in uploaded documents (e.g., "summarize document X", "what does the PDF say about Y?", "details about topic Z from the file", "tell me more about X mentioned in the document").
- "direct_answer": Use this for greetings, general knowledge questions, simple conversational follow-ups not requiring document lookup, or commands.

Here are some examples of how to decide:
---
User Question: "Hello there!"
Your decision: direct_answer
---
User Question: "Can you summarize the main points of the uploaded PDF about robot navigation?"
Your decision: rag_search
---
User Question: "What is the capital of Spain?"
Your decision: direct_answer
---
User Question: "Tell me more about the limitations of the controller discussed in the document."
Your decision: rag_search
---

Which action should be taken for the current User Question? Respond with ONLY the word "rag_search" or "direct_answer". Do not add any other text.
Your decision:"""
    
    router_call_parts = [{"role": "user", "parts": [router_prompt]}]
    
    try:
        decision_text_raw = generate_with_gemini(GEMINI_ROUTER_MODEL_NAME, router_call_parts, temperature=0.1, max_output_tokens=10, is_router_call=True)
        decision_text_clean = decision_text_raw.lower().replace("action:", "").replace("decision:", "").strip()
        decision_text_clean = ''.join(filter(str.isalnum, decision_text_clean))
        print(f"Gemini Router Raw Decision Text: '{decision_text_raw}' -> Cleaned: '{decision_text_clean}'")
        if "ragsearch" in decision_text_clean: action_to_take = "rag_search"
        elif "directanswer" in decision_text_clean: action_to_take = "direct_answer"
        else:
            print(f"Gemini Router unclear: '{decision_text_raw}'. Fallback heuristic.")
            rag_keywords = ["document", "pdf", "file", "summary", "summarize", "detail", "information on", "content of", "what does it say about"]
            if any(keyword in user_message_text.lower() for keyword in rag_keywords): action_to_take = "rag_search"; print("Fallback: RAG.")
            else: action_to_take = "direct_answer"; print("Fallback: Direct.")
    except Exception as e:
        print(f"Gemini Router call Error: {e}. Fallback heuristic.")
        rag_keywords = ["document", "pdf", "file", "summary", "summarize", "detail", "information on", "content of", "what does it say about"]
        if any(keyword in user_message_text.lower() for keyword in rag_keywords): action_to_take = "rag_search"; print("Fallback (error): RAG.")
        else: action_to_take = "direct_answer"; print("Fallback (error): Direct.")

    print(f"Final Router decision: {action_to_take.upper()} for: '{user_message_text}'")

    context_for_llm_prompt_str = "" 
    retrieved_sources_details = []

    if action_to_take == "rag_search":
        try:
            query_embedding = embedding_model.encode(user_message_text).tolist()
            initial_hits = qdrant_client.search(collection_name=QDRANT_COLLECTION_NAME, query_vector=query_embedding, limit=INITIAL_QDRANT_SEARCH_LIMIT)
            final_search_results = []
            if initial_hits:
                if cross_encoder_model:
                    sentence_pairs = [[user_message_text, hit.payload['text']] for hit in initial_hits if hit.payload and 'text' in hit.payload]
                    if sentence_pairs:
                        scores = cross_encoder_model.predict(sentence_pairs)
                        scored_hits = sorted(zip(scores, initial_hits), key=lambda x: x[0], reverse=True)
                        final_search_results = [hit for score, hit in scored_hits[:FINAL_CONTEXT_CHUNKS_LIMIT]]
                else: final_search_results = initial_hits[:FINAL_CONTEXT_CHUNKS_LIMIT]
            
            if final_search_results:
                context_parts_text = []
                temp_char_count = 0 # Using char count as a proxy for tokens for Gemini context limit
                for i, hit in enumerate(final_search_results):
                    chunk_text = hit.payload['text']
                    if (temp_char_count + len(chunk_text)) > (MAX_CONTEXT_TOKENS_FOR_LLM * 3): # Approx 3 chars/token
                        print(f"RAG: Approx char limit ({MAX_CONTEXT_TOKENS_FOR_LLM * 3}) reached for context. Stopping.")
                        break
                    context_parts_text.append(f"Source Document: '{hit.payload['source']}' (Context Snippet {i+1}):\n{chunk_text}")
                    retrieved_sources_details.append({"filename": hit.payload['source'], "id": str(hit.id), "text": chunk_text})
                    temp_char_count += len(chunk_text)
                
                if context_parts_text: # MODIFIED PROMPT for no in-text chunk citation
                    context_for_llm_prompt_str = ("Based ONLY on the following context from uploaded documents, answer the user's question. "
                                               "Do NOT mention specific context snippet numbers or chunk IDs in your answer. "
                                               "If the answer isn't in the context, state that the information is not found in the provided documents.\n\n"
                                               "Provided Context:\n" + "\n---\n".join(context_parts_text) + "\n---\n")
                    print(f"RAG: Prepared context with {len(context_parts_text)} snippets.")
                else:
                    action_to_take = "direct_answer"; retrieved_sources_details = []; context_for_llm_prompt_str = ""
                    print("RAG: No usable snippets after filtering, fallback to direct_answer.")
            else:
                action_to_take = "direct_answer"; retrieved_sources_details = []; context_for_llm_prompt_str = ""
                print("RAG: No initial matches, fallback to direct_answer.")
        except Exception as e:
            print(f"RAG Processing Error: {e}. Fallback to direct_answer.")
            action_to_take = "direct_answer"; retrieved_sources_details = []; context_for_llm_prompt_str = ""

    system_instruction_text = "You are a helpful and concise AI assistant."
    # Specific RAG instruction is now part of context_for_llm_prompt_str
    
    current_user_turn_content_parts = []
    if action_to_take == "rag_search" and context_for_llm_prompt_str:
        # Prepend RAG context and instructions, then the user question
        current_user_turn_content_parts.append(context_for_llm_prompt_str + f"\nUser Question: {user_message_text}")
    else:
        # For direct answer, or if RAG failed and context_for_llm_prompt_str is empty
        # Prepend system instruction if no history, otherwise just the user message
        if not conversation_history:
             current_user_turn_content_parts.append(f"{system_instruction_text}\n\n{user_message_text}")
        else:
            current_user_turn_content_parts.append(user_message_text)

    gemini_chat_history = [msg for msg in conversation_history] # Copy existing history
    gemini_chat_history.append({"role": "user", "parts": current_user_turn_content_parts})
    
    response_text = generate_with_gemini(GEMINI_CHAT_MODEL_NAME, gemini_chat_history, temperature=0.6, max_output_tokens=1500)
    print(f"Gemini Final Raw Response: {response_text}")

    conversation_history.append({"role": "user", "parts": [user_message_text]}) 
    conversation_history.append({"role": "model", "parts": [response_text]}) 
    
    final_sources_to_return = retrieved_sources_details if action_to_take == "rag_search" and retrieved_sources_details else []
    return jsonify({"response": response_text, "sources": final_sources_to_return})

@app.route('/upload', methods=['POST'])
def upload_file():
    global text_splitter
    if 'files[]' not in request.files: return jsonify({"error": "No file part"}), 400
    files = request.files.getlist('files[]')
    if not files or all(not f.filename for f in files): return jsonify({"error": "No files selected or files have no names"}), 400

    results = []
    saved_file_paths = {}

    for file_obj in files:
        original_filename = file_obj.filename
        if not original_filename: continue
        file_result = {"original_filename": original_filename, "status": "error", "message": "Skipped"}
        file_path = None
        if allowed_file(original_filename):
            secured_filename_for_save = secure_filename(original_filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], secured_filename_for_save)
            saved_file_paths[original_filename] = file_path 
            try:
                file_obj.save(file_path)
                text_content = ""
                if original_filename.lower().endswith('.pdf'):
                    with fitz.open(file_path) as doc_pdf:
                        page_texts = [p.get_text("text", sort=True).strip() for p in doc_pdf if p.get_text("text")]
                    text_content = "\n\n".join(page_texts)
                elif original_filename.lower().endswith('.docx'):
                    doc_docx = docx.Document(file_path) 
                    text_content = "\n".join([p.text for p in doc_docx.paragraphs if p.text.strip()])
                elif original_filename.lower().endswith(('.txt', '.md')):
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f: text_content = f.read()
                    except UnicodeDecodeError:
                        with open(file_path, 'r', encoding='latin-1') as f: text_content = f.read()
                if not text_content.strip(): raise ValueError("No text extracted.")
                langchain_docs_for_splitting = [LangchainDocument(page_content=text_content, metadata={"source": original_filename})]
                chunks_objects = text_splitter.split_documents(langchain_docs_for_splitting)
                chunks = [doc.page_content for doc in chunks_objects]
                if not chunks: raise ValueError("No chunks generated.")
                embeddings = embedding_model.encode(chunks, show_progress_bar=False, batch_size=32).tolist()
                points = [models.PointStruct(id=str(uuid.uuid4()), vector=emb, payload={"text": chk, "source": original_filename}) for emb, chk in zip(embeddings, chunks)]
                qdrant_client.upsert(collection_name=QDRANT_COLLECTION_NAME, points=points, wait=True)
                file_result["status"] = "success"; file_result["message"] = f"Processed, {len(chunks)} chunks."
            except Exception as e:
                print(f"ERROR processing '{original_filename}': {e}")
                file_result["status"] = "error"; file_result["message"] = f"Failed: {str(e)}"
        else:
            file_result["status"] = "error"; file_result["message"] = "File type not allowed"
        results.append(file_result)

    for f_path_to_clean in saved_file_paths.values():
        if os.path.exists(f_path_to_clean):
            try: os.remove(f_path_to_clean)
            except Exception as del_e: print(f"WARN: Could not remove '{f_path_to_clean}': {del_e}")
    if not results: return jsonify({"error": "No valid files to process."}), 400
    all_successful = all(r['status'] == 'success' for r in results)
    any_successful = any(r['status'] == 'success' for r in results)
    status_code = 200
    if not any_successful: status_code = 400
    elif not all_successful: status_code = 207 
    return jsonify({"results": results}), status_code

if __name__ == "__main__":
    print(f"\n--- Starting Flask Server (Gemini Mode) ---")
    if not GOOGLE_API_KEY:
        print("FATAL: GOOGLE_API_KEY is not set. Please set it in a .env file or as an environment variable.")
    else:
        print("Gemini API Key detected.")
    print("\n--- Flask App Ready ---")
    app.run(host="0.0.0.0", port=5001, debug=False)